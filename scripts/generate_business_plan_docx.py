from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/dingsen/Documents/echoland/Echoland回声地商业计划书.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 24, 31)
MUTED = RGBColor(92, 100, 112)
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "E8EEF5"
GREEN = RGBColor(168, 255, 61)
GREEN_HEX = "A8FF3D"
BLACK = RGBColor(0, 0, 0)


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9DEE7", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def format_table(table, widths=None, header=True):
    set_table_borders(table)
    if widths:
        set_table_widths(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            if header and row_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)


def add_para(doc, text="", style=None, bold=False, color=INK, size=11, align=None, after=8):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 1:
            set_run_font(run, size=16, bold=True, color=BLUE)
        elif level == 2:
            set_run_font(run, size=13, bold=True, color=BLUE)
        else:
            set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    format_table(table, widths=[6.3], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image_placeholder(doc, title, suggestion):
    table = doc.add_table(rows=1, cols=1)
    format_table(table, widths=[6.3], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"[图片占位] {title}")
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(f"建议插入：{suggestion}")
    set_run_font(r2, size=9.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_source(doc, label, url):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(f"{label}：{url}")
    set_run_font(r, size=7.8, color=MUTED)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    format_table(table, widths=widths, header=True)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=9.5, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9.2, color=INK)
    format_table(table, widths=widths, header=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Echoland 回声地商业计划书 | 内部讨论稿"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential - Course / Venture Planning Draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=8.5, color=MUTED)


def build_doc():
    doc = Document()
    configure_doc(doc)

    # Cover
    add_para(doc, "ECHOLAND 回声地", bold=True, color=BLACK, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "商业计划书", bold=True, color=BLUE, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "有线下出口的氛围型城市社交社区", color=MUTED, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, "版本：内部讨论稿 | 日期：2026 年 6 月 | 首发城市建议：南京", color=MUTED, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_image_placeholder(
        doc,
        "封面主视觉",
        "一张 3-4 个手机界面叠加南京夜间城市背景的品牌视觉图；画面应包含白色 App 界面、荧光绿强调色、地图 marker、社区短帖和 AI 路线卡片。",
    )
    add_callout(
        doc,
        "核心判断",
        "Echoland 的机会不在于做一个更全的活动平台，而在于把社交氛围、城市小众体验和线下组织力合在一起。用户可以因为社区 vibe 留下，也可以因为想打破无聊生活而出门。商业化来自真实到场、社区影响力和线下文化 IP。",
    )
    doc.add_page_break()

    # Contents
    add_heading(doc, "目录", 1)
    contents = [
        "1. 执行摘要",
        "2. 项目定义与核心假设",
        "3. 问题定义：年轻人的无聊生活与线下社交断层",
        "4. 市场与场景分析",
        "5. 目标用户与关键使用场景",
        "6. 产品、服务与业务体系",
        "7. 内容供给体系：PGC 活动与 UGC 回声点",
        "8. 技术实现路径与 AI 应用点",
        "9. 创新点与核心壁垒",
        "10. 竞品分析",
        "11. 商业模式与收入设计",
        "12. 运营与增长策略",
        "13. 风险、商业漏洞与应对方案",
        "14. 阶段里程碑与关键指标",
        "15. 图片插入清单",
        "16. 资料来源",
    ]
    add_numbered(doc, contents)
    doc.add_page_break()

    add_heading(doc, "1. 执行摘要", 1)
    add_para(
        doc,
        "Echoland 回声地是一款面向年轻人的城市小众活动发现、AI 路线规划与氛围型社交 App。它面向 18-30 岁城市青年，帮助他们摆脱重复的聚餐、商场、电影和无意识刷屏，找到有趣的活动、微小的城市点位、同频的人和可执行的线下路线。",
    )
    add_para(
        doc,
        "产品由四个核心模块组成：Threads 风格的社区信息流、PGC 小众活动、UGC 回声点、AI 路线与邀约。社交是核心，不是附属模块。Echoland 希望形成一种独特的社区气质，让用户即使暂时不出门，也愿意因为这里的人、表达方式和城市审美而日常打开；活动、回声点与路线则让这种线上气质拥有真实的线下出口。",
    )
    add_simple_table(
        doc,
        ["维度", "核心结论"],
        [
            ["用户价值", "找到有趣的地方、有趣的活动和有趣的人，降低出门决策成本与独自参加活动的心理门槛。"],
            ["市场机会", "服务消费、现场音乐、小型活动和年轻人线下社交需求同时增长，存在面向城市文化生活的细分机会。"],
            ["产品差异", "不做大众点评式商户评价，不做小红书式瀑布流种草，不做单纯票务平台，而是做有线下出口的氛围型城市社交社区。"],
            ["商业化路径", "早期通过 club/livehouse/场地联名与核销分成验证价值，中期做社区原生广告和赞助路线，后期发展 Echoland Night、城市周末和音乐节类 IP。"],
            ["关键风险", "社区 vibe 难以生产、UGC 安全与质量控制、场地合作冷启动、城市扩张重运营、广告商业化伤害社区调性。"],
        ],
        [1.2, 5.1],
    )
    add_image_placeholder(
        doc,
        "Echoland 商业逻辑总览图",
        "一张四象限或飞轮图，展示“社区 vibe -> 城市内容 -> AI 路线/邀约 -> 线下体验 -> 回流社区 -> 商业化”的循环。",
    )

    add_heading(doc, "2. 项目定义与核心假设", 1)
    add_heading(doc, "2.1 产品定义", 2)
    add_para(
        doc,
        "Echoland 的产品定义是：有线下出口的氛围型城市社交社区。它提供小众活动、微小城市点位和 AI 路线，但这些功能服务于一个更大的目标：让年轻人从无聊、重复、低质量的城市生活里逃出来，进入一个更有趣、更同频、更能真实见面的城市生活网络。",
    )
    add_heading(doc, "2.2 不做什么", 2)
    add_bullets(
        doc,
        [
            "不做大众点评式的全量商户评价平台。Echoland 不追求覆盖所有餐厅、商场和景点。",
            "不做美团式交易平台。早期不以补贴、团购、支付闭环为核心。",
            "不做小红书瀑布流。社区以短帖和同频表达为主，不把商业化内容变成密集种草。",
            "不做纯票务平台。票务或报名只是线下转化工具，不是产品的唯一价值。",
            "不做普通 AI 行程助手。AI 的作用是压缩选择、生成可执行路线和帮助用户迈出门。",
        ],
    )
    add_heading(doc, "2.3 关键假设", 2)
    add_bullets(
        doc,
        [
            "南京等城市存在一批 18-30 岁年轻人，他们厌倦常规消费，愿意寻找更小众、更有氛围的城市体验。",
            "用户不仅会因为活动而进入，也会因为社区氛围、表达方式和人群气质而留下。",
            "PGC 活动能提供可执行和可商业化的供给，UGC 回声点能提供城市发现感和内容壁垒。",
            "AI 路线可以降低决策成本，把浏览转化成真实出门。",
            "当平台能证明到场能力后，club、livehouse、书店、放映空间和市集主办方愿意进行分成或联合活动。",
        ],
    )

    add_heading(doc, "3. 问题定义：年轻人的无聊生活与线下社交断层", 1)
    add_para(
        doc,
        "当代年轻人并不缺信息。他们缺的是能让自己行动起来的信息。周末晚上刷小红书、短视频和朋友圈时，用户经常能看到大量生活灵感，但这些灵感常常分散、重复、难判断，也缺少清晰的出门路径。更深的问题是，城市生活正在变得可消费却不一定有趣。聚餐、商场、电影和热门景点能填满时间，却很难带来真正的新鲜感与身份认同。",
    )
    add_para(
        doc,
        "社交层面同样存在断层。线上社区能提供情绪共鸣和表达空间，但关系停在线上；线下活动能提供真实连接，但发现成本高、独自参加有心理门槛。Echoland 的问题定义不是“用户找不到活动”，而是“用户不知道怎样从无聊生活中找到一个值得出门的理由，并遇到同频的人”。",
    )
    add_simple_table(
        doc,
        ["用户痛点", "具体表现", "Echoland 的回应"],
        [
            ["生活重复", "下班和周末的选择被聚餐、商场、电影占据。", "提供小众活动、回声点和主题路线。"],
            ["信息分散", "活动散落在公众号、小红书、微信群、主办方账号和线下海报。", "通过运营、场地投稿和 AI 提取进行结构化。"],
            ["选择困难", "活动太多，不知道哪个适合自己的预算、时间和兴趣。", "AI 根据时间、预算、兴趣和氛围生成路线。"],
            ["独自门槛", "一个人去 livehouse、放映、club 或夜游容易尴尬。", "在详情、路线和社区中提供轻邀约入口。"],
            ["线上关系弱", "社区有 vibe，但缺少线下凝结点。", "用活动、路线、联名夜和线下 IP 把人聚到同一现场。"],
        ],
        [1.1, 2.45, 2.75],
    )
    add_image_placeholder(
        doc,
        "用户问题旅程图",
        "从“无聊刷手机”到“发现活动/回声点”再到“AI 生成路线、发起邀约、线下见面、回社区发帖”的用户旅程图。",
    )

    add_heading(doc, "4. 市场与场景分析", 1)
    add_heading(doc, "4.1 服务消费与快乐生活场景增长", 2)
    add_para(
        doc,
        "美团研究院在 2025 年生活服务消费趋势中测算，我国服务零售市场规模整体接近 30 万亿元，其中覆盖休闲娱乐、运动健身、丽人、宠物等行业的“快乐生活”消费市场规模超过 3 万亿元。该报告也提出“热爱当下，重返线下”的消费趋势，说明年轻人愿意为即时体验、情绪价值和真实连接付费。[资料：美团研究院，2026-01]",
    )
    add_heading(doc, "4.2 现场演出与小型音乐活动提供商业化抓手", 2)
    add_para(
        doc,
        "中国演出行业协会数据显示，2025 年全国营业性演出场次 64.04 万场，票房收入 616.55 亿元，观众人数 1.94 亿人次。中演协与灯塔发布的现场音乐报告披露，2025 年 5000 人以下中小型现场音乐演出票房达 23.15 亿元，场次达 4.41 万场，观演人次达 886.73 万人。这类活动规模更小、圈层更明确、和 Echoland 的社区气质更匹配。[资料：人民日报，2026-01；新浪财经，2026-04]",
    )
    add_heading(doc, "4.3 年轻人社交正在回到具体场景", 2)
    add_para(
        doc,
        "复旦发展研究院发布的 2025 年轻人生活方式报告显示，26-30 岁受访者在线下活动参与方面占比较高，工作与兴趣成为当代年轻人寻找“搭子”的重要连接媒介。对 Echoland 来说，社交不应停留在聊天和关注关系，而应围绕真实的地点、活动和路线发生。[资料：复旦发展研究院，2025]",
    )
    add_simple_table(
        doc,
        ["市场层面", "数据或趋势", "对 Echoland 的意义"],
        [
            ["服务消费", "服务零售整体接近 30 万亿元，“快乐生活”消费超过 3 万亿元。", "体验消费具备大盘支撑，城市小众活动是细分切口。"],
            ["演出市场", "2025 年营业性演出票房 616.55 亿元，观众 1.94 亿人次。", "用户为线下现场付费的习惯已经存在。"],
            ["中小型现场音乐", "5000 人以下演出票房 23.15 亿元，4.41 万场。", "适合从 club、livehouse、小剧场、独立空间切入。"],
            ["年轻人社交", "兴趣与线下活动成为“搭子”连接的重要媒介。", "社区与线下出门场景应绑定设计。"],
        ],
        [1.2, 2.35, 2.75],
    )
    add_image_placeholder(
        doc,
        "市场规模与场景机会图",
        "柱状图或漏斗图：服务零售 30 万亿元 -> 快乐生活 3 万亿元 -> 演出 616.55 亿元 -> 中小型现场音乐 23.15 亿元 -> Echoland 首发城市可切入场景。",
    )

    add_heading(doc, "5. 目标用户与关键使用场景", 1)
    add_para(
        doc,
        "Echoland 的目标用户不是所有本地生活消费者，而是对城市生活有审美要求、对常规娱乐感到疲惫、愿意在真实世界里寻找新鲜体验和同频关系的年轻人。用户可以因为想玩点不一样的而来，也可以因为喜欢社区氛围而来。",
    )
    add_simple_table(
        doc,
        ["用户分层", "画像", "核心需求", "关键功能"],
        [
            ["Vibe 型社交用户", "18-30 岁，喜欢短句表达、城市审美和同频互动。", "日常打开一个有气质的社区。", "Threads 风信息流、关注、评论、发布短帖。"],
            ["反无聊探索者", "学生或初入职场者，厌倦常规聚餐和商场。", "找到今晚或周末的新鲜去处。", "地图发现、精选活动、回声点、主题筛选。"],
            ["小众文化爱好者", "关注独立音乐、放映、展览、市集、摄影、城市漫游。", "发现更垂直、更有调性的活动。", "PGC 活动、场地主页、活动详情、收藏。"],
            ["轻社交/搭子用户", "想参加线下活动，但一个人去有心理门槛。", "低压力找到同频同行者。", "发起邀约、加入邀约、路线节点邀约。"],
            ["新城市居民", "刚到南京或新城市，对文化空间不熟。", "快速找到本地有趣生活半径。", "AI 路线、城市专题、回声点地图。"],
        ],
        [1.15, 1.35, 1.75, 2.05],
    )
    add_heading(doc, "5.1 高频场景", 2)
    add_bullets(
        doc,
        [
            "周五晚上：用户想找一个不只是吃饭的去处，打开 Echoland 浏览今晚活动和附近回声点。",
            "周末下午：用户输入预算、时长、兴趣，让 AI 生成一条放映 + 书店 + 黑胶酒吧路线。",
            "日常碎片时间：用户刷社区短帖，看到同频用户分享夜游、livehouse 和隐秘点位。",
            "陌生城市生活：新城市居民通过 Echoland 找到本地文化空间，并通过轻邀约认识人。",
            "线下活动之后：用户回到社区发布照片、现场感受和新发现的回声点。",
        ],
    )

    add_heading(doc, "6. 产品、服务与业务体系", 1)
    add_heading(doc, "6.1 产品架构", 2)
    add_para(
        doc,
        "产品前台采用四个底部 Tab：首页、AI 路线、社区、我的。邀约不是独立 Tab，而是贯穿在活动详情、路线节点、社区帖子和我的页面中的轻社交能力。这样的结构可以避免产品显得过重，同时让线下见面自然嵌入各个使用路径。",
    )
    add_simple_table(
        doc,
        ["模块", "用户价值", "核心内容", "商业价值"],
        [
            ["首页", "快速发现城市里有趣的活动和回声点。", "城市、搜索、标签筛选、地图 marker、精选活动卡。", "承载活动曝光、场地合作和推广位。"],
            ["AI 路线", "把想法变成今晚或周末可执行计划。", "预算、时长、兴趣、氛围、路线结果、加入路线。", "提升到场转化，为场地分成提供数据。"],
            ["社区", "形成日常打开和同频表达。", "短帖、图片、关联活动、点赞评论分享、发起邀约。", "承载原生广告、品牌赞助和社区影响力。"],
            ["我的", "沉淀用户行为记录。", "收藏、路线、邀约、帖子、等级、偏好。", "支撑会员、权益和复访。"],
            ["邀约", "降低一个人参加线下活动的门槛。", "集合时间、地点、人数、备注、状态。", "把线上兴趣转成真实到场和组织力。"],
        ],
        [0.95, 1.55, 2.1, 1.7],
    )
    add_image_placeholder(
        doc,
        "产品信息架构图",
        "一张 App 架构图，展示四个 Tab、活动详情子页面、邀约 Bottom Sheet，以及活动/回声点/路线/社区帖子之间的跳转关系。",
    )
    add_heading(doc, "6.2 业务体系", 2)
    add_para(
        doc,
        "Echoland 的业务不只是 App。它至少包含三套业务：线上社区运营、城市内容供给、线下场地合作。线上社区负责日常粘性，城市内容负责打破无聊，线下合作负责商业化和真实到场。",
    )
    add_simple_table(
        doc,
        ["业务线", "主要工作", "早期目标"],
        [
            ["社区运营", "种子用户筛选、社区调性维护、话题与内容精选、UGC 审核。", "形成一批用户愿意日常打开和发帖的社区氛围。"],
            ["城市内容", "PGC 活动录入、UGC 回声点审核、主题路线策划、AI 路线素材库。", "在南京构建足够密度的活动与点位供给。"],
            ["场地合作", "访谈老板、谈联名夜、核销分成、活动推广和数据反馈。", "用小场地证明 Echoland 能带来到场用户。"],
            ["线下 IP", "Echoland Night、城市周末、音乐/放映/市集联动活动。", "验证社区组织力，积累品牌影响力。"],
        ],
        [1.2, 3.15, 1.95],
    )

    add_heading(doc, "7. 内容供给体系：PGC 活动与 UGC 回声点", 1)
    add_para(
        doc,
        "Echoland 的内容供给由 PGC 与 UGC 两条线组成。PGC 是正式活动供给，主要来自场地老板、主办方和运营编辑。UGC 是用户发现的微小、有趣、不为人知的城市点位，被称为“回声点”。两者共同形成城市体验内容库：PGC 让用户有明确可去的活动，UGC 让产品拥有小众感、发现感和社区生命力。",
    )
    add_simple_table(
        doc,
        ["类型", "来源", "内容样例", "平台价值"],
        [
            ["PGC 活动", "club、livehouse、书店、影像空间、市集主办方、城市编辑。", "独立音乐、放映、市集、展览、书店沙龙、夜游、手作工作坊。", "可执行、可核销、可商业化。"],
            ["UGC 回声点", "普通用户、城市观察者、摄影爱好者、社区 KOC。", "雨后倒影墙、夜景天桥、旧厂房角落、隐秘小店、河边发呆点。", "小众感、社区内容资产、路线灵魂。"],
            ["编辑精选", "运营团队从 PGC 和 UGC 中筛选。", "本周 10 个夜游回声点、周五黑胶路线、独处友好地图。", "建立审美筛选权和品牌调性。"],
            ["AI 路线素材", "活动、回声点、场地、用户偏好和距离预算。", "3 小时 100 元以内的南京夜晚路线。", "把内容转化成行动。"],
        ],
        [1.0, 1.4, 2.1, 1.8],
    )
    add_heading(doc, "7.1 场地端供给方式", 2)
    add_bullets(
        doc,
        [
            "早期：运营团队人工关注公众号、小红书、微信群、主办方账号和线下海报，录入和编辑活动。",
            "中期：为场地老板提供轻量投稿表单或小程序，支持粘贴推文链接、上传海报、填写时间地点价格。",
            "自动化：用 OCR 和 AI 抽取活动名称、时间、地点、价格、标签和简介，由人工审核后发布。",
            "数据反馈：给场地展示曝光、收藏、路线加入、报名、核销和邀约数据，让老板看到增量价值。",
        ],
    )
    add_heading(doc, "7.2 UGC 回声点机制", 2)
    add_bullets(
        doc,
        [
            "上传字段应保持轻量：位置、1-3 张照片、一句话描述、标签、最佳时间、注意事项。",
            "UGC 需要审核：过滤危险地点、私密住宅、违法内容、恶意曝光和侵犯隐私的点位。",
            "社区反馈应参与质量控制：用户可标记“真的有趣”“已失效”“不安全”“位置不准”。",
            "编辑精选负责把优质回声点变成首页推荐、主题榜单和 AI 路线节点。",
        ],
    )
    add_image_placeholder(
        doc,
        "PGC + UGC 内容供给流程图",
        "左侧展示场地老板/运营录入 PGC 活动，右侧展示用户上传 UGC 回声点，中间经过 AI 提取、审核、编辑精选，最终进入首页、社区和 AI 路线。",
    )

    add_heading(doc, "8. 技术实现路径与 AI 应用点", 1)
    add_heading(doc, "8.1 当前 Demo 技术路径", 2)
    add_para(
        doc,
        "课程展示阶段的技术目标是高保真前端 Demo，不追求真实后端、数据库、登录或真实 AI API。建议继续采用单文件 HTML App、本地假数据、前端状态模拟和可降级地图方案。重点是演示路径稳定、视觉统一、交互完整。",
    )
    add_simple_table(
        doc,
        ["层级", "课程 Demo 做法", "未来产品化升级"],
        [
            ["前端", "HTML / CSS / JavaScript 单文件输出，四个 Tab 和本地状态。", "React Native / Flutter 或 WebView Hybrid，支持跨端发布。"],
            ["数据", "本地 JSON 假数据：活动、回声点、帖子、路线、邀约。", "活动内容 CMS、UGC 数据库、场地端投稿系统。"],
            ["地图", "固定南京中心，固定 marker，必要时使用地图截图降级。", "真实地图 SDK、地理围栏、路径距离估算、POI 增强。"],
            ["AI", "预设路线模拟 1 秒 Loading。", "活动抽取、标签生成、路线生成、内容审核辅助、个性化推荐。"],
            ["社交", "点赞、评论、分享、邀约均前端模拟。", "真实账号体系、关系链、消息、风控和安全机制。"],
        ],
        [1.0, 2.65, 2.65],
    )
    add_heading(doc, "8.2 AI 应用点", 2)
    add_para(
        doc,
        "AI 在 Echoland 中不应被包装成噱头。它的角色是压缩选择成本、提高运营效率和帮助用户出门。最有价值的 AI 应用不在于聊天，而在于将分散、非结构化、带有审美差异的信息转化成可执行的路线和可运营的数据。",
    )
    add_simple_table(
        doc,
        ["AI 应用点", "输入", "输出", "价值"],
        [
            ["活动信息抽取", "公众号推文、海报截图、主办方文案。", "标题、时间、地点、价格、类型、标签、简介。", "降低运营录入成本。"],
            ["回声点标签生成", "用户上传图片和一句话描述。", "适合独处、夜晚更好、适合拍照、免费、安静等标签。", "提升 UGC 可检索性。"],
            ["路线生成", "时间、预算、兴趣、氛围、位置、活动和回声点库。", "2-4 个节点的可执行路线。", "把浏览转成行动。"],
            ["推荐理由生成", "活动详情、标签、用户偏好。", "为什么适合你、适合谁、不适合谁。", "帮助判断是否值得去。"],
            ["内容审核辅助", "UGC 文本、图片、位置。", "风险标签、重复检测、疑似危险点位提示。", "降低安全和内容风险。"],
            ["运营洞察", "点击、收藏、加入路线、邀约、到场数据。", "活动热度、场地效果、用户偏好。", "服务商业化和场地合作。"],
        ],
        [1.25, 1.55, 1.95, 1.55],
    )
    add_image_placeholder(
        doc,
        "AI 应用架构图",
        "展示“内容输入 -> AI 抽取/打标/审核 -> 活动与回声点数据库 -> 路线生成 -> 用户反馈数据回流”的技术流程图。",
    )

    add_heading(doc, "9. 创新点与核心壁垒", 1)
    add_heading(doc, "9.1 创新点", 2)
    add_bullets(
        doc,
        [
            "把氛围型社交与真实线下去处绑定，而不是停留在纯线上社区。",
            "把正式活动与微小城市点位并列为体验内容，让城市发现从“商户/景点”扩展为“回声点”。",
            "用 AI 路线把兴趣、时间、预算和地点串成行动，降低用户出门前的决策成本。",
            "用轻邀约把同频感转化为低压力的线下见面。",
            "用场地分成、赞助路线和自有活动 IP 让社区影响力拥有清晰商业出口。",
        ],
    )
    add_heading(doc, "9.2 可能形成的壁垒", 2)
    add_simple_table(
        doc,
        ["壁垒", "形成机制", "复制难点"],
        [
            ["社区气质", "首批用户、内容规范、编辑精选和视觉语言共同塑造。", "竞品可以抄功能，但很难直接复制人群和表达方式。"],
            ["UGC 回声点数据库", "用户持续上传非标准 POI 和微小体验点。", "这些内容不在传统商户库里，需要社区参与和审核。"],
            ["场地关系网", "与 club、livehouse、书店、放映空间形成合作和联名。", "依赖本地运营、信任和长期到场数据。"],
            ["线下组织力", "平台能让线上用户真实聚到同一现场。", "这是比内容流量更难复制的能力。"],
            ["AI 路线反馈数据", "用户收藏、加入路线、发起邀约、到场和发帖形成闭环。", "数据越多，路线越符合社区人群偏好。"],
        ],
        [1.1, 2.6, 2.6],
    )

    add_heading(doc, "10. 竞品分析", 1)
    add_para(
        doc,
        "Echoland 的竞品并不只是一类产品。它的竞争来自内容社区、票务平台、本地生活平台、地图平台和陌生人社交。每类竞品都能复制部分功能，但都存在短板。Echoland 必须守住自己的主轴：社交氛围、城市内容和线下组织力三者的结合。",
    )
    add_simple_table(
        doc,
        ["竞品", "强项", "短板", "Echoland 切入点"],
        [
            ["小红书", "生活方式内容、搜索心智、商业化能力强。", "信息过载，线下组织和即时路线弱。", "把灵感变成路线和邀约，做更轻的同频社交。"],
            ["Threads", "轻表达、关注关系、社区 vibe。", "本地活动和线下场景弱。", "将氛围型社交导向城市线下体验。"],
            ["豆瓣同城", "小众活动和兴趣社区心智接近。", "产品老化、移动端体验弱、年轻化不足。", "更年轻的视觉、AI 路线和即时互动。"],
            ["秀动 ShowStart", "独立音乐票务、Livehouse 场景、演出沟通群。", "社区日常粘性和非演出城市点位弱。", "覆盖更广城市体验，将演出与回声点、路线和社交结合。"],
            ["大麦", "大型演出、票务交易和供给规模。", "偏交易，缺少小众社区和活动前后关系。", "从小场地、club、独立空间切入。"],
            ["美团/大众点评/高德", "商户、地图、交易、评价和流量强。", "偏大众消费，非商业点位和社区氛围弱。", "做非标准、小众、有审美的城市体验。"],
            ["Soul/陌陌类社交", "陌生人社交和关系匹配。", "容易缺少真实见面理由和场景质量。", "用活动、路线和回声点提供低尴尬见面理由。"],
        ],
        [1.05, 1.65, 1.65, 1.95],
    )
    add_image_placeholder(
        doc,
        "竞品定位图",
        "二维坐标图：横轴为“工具/交易 -> 社交/社区”，纵轴为“大众消费 -> 小众文化体验”。把小红书、豆瓣同城、秀动、大麦、美团、高德、Threads 和 Echoland 放入图中。",
    )

    add_heading(doc, "11. 商业模式与收入设计", 1)
    add_heading(doc, "11.1 商业模式画布", 2)
    add_simple_table(
        doc,
        ["模块", "设计"],
        [
            ["客户细分", "C 端为城市青年、小众文化爱好者、轻社交用户、新城市居民；B 端为 club、livehouse、书店、放映空间、市集主办方和生活方式品牌。"],
            ["价值主张", "帮助用户找到有趣的地方、有趣的活动和有趣的人；帮助场地获得精准到场用户；帮助品牌进入高粘性城市青年社区。"],
            ["渠道", "小红书/抖音内容种草、校园社群、KOC、场地联名、线下活动、朋友邀约传播。"],
            ["客户关系", "关注关系、邀约关系、收藏路线、城市身份标签、活动参与记录。"],
            ["收入来源", "场地分成、活动推广、社区原生广告、赞助路线、品牌联名、Echoland Night、城市文化 IP。"],
            ["核心资源", "社区气质、PGC 活动库、UGC 回声点、场地关系网、城市编辑能力、到场数据。"],
            ["关键业务", "社区运营、内容审核、场地 BD、AI 路线运营、活动组织和安全风控。"],
            ["重要伙伴", "club、livehouse、独立书店、影像空间、市集主办方、摄影/音乐/艺术 KOC、品牌赞助方。"],
            ["成本结构", "运营编辑、社区审核、BD、活动执行、技术开发、市场获客和安全机制。"],
        ],
        [1.45, 4.85],
    )
    add_heading(doc, "11.2 收入路径", 2)
    add_simple_table(
        doc,
        ["阶段", "收入方式", "说明"],
        [
            ["第一阶段：证明到场", "club/livehouse 分成、核销佣金、联名夜分账。", "先用小场地验证平台能否带来真实到场用户。"],
            ["第二阶段：证明人群价值", "社区原生广告、赞助路线、活动 Boost。", "面向相机、香水、精酿、唱片、服饰、咖啡等生活方式品牌。"],
            ["第三阶段：证明 IP", "Echoland Night、Echoland Weekend、城市文化节、音乐节。", "当社区影响力形成后，将线上组织力转化为自有线下 IP。"],
        ],
        [1.45, 1.8, 3.05],
    )
    add_heading(doc, "11.3 Club 与场地分成模型", 2)
    add_bullets(
        doc,
        [
            "票务分成：用户通过 Echoland 报名或购票，平台抽取 5%-15%。",
            "核销佣金：用户到场出示 Echoland 码，场地按有效到场人数支付固定费用。",
            "酒水套餐分成：Echoland 设计联名套餐，按订单或毛利分成。",
            "联名活动分账：平台负责主题策划、用户招募和内容传播，场地负责空间、基础服务和现场执行。",
        ],
    )
    add_heading(doc, "11.4 社区广告原则", 2)
    add_para(
        doc,
        "社区广告必须克制。Echoland 的广告不能像促销信息流，更适合做 sponsored post、赞助路线、城市专题和品牌联名活动。广告主应与城市生活方式和小众文化场景相关，品牌应融入用户想要参与的体验，而不是打断社区内容。",
    )
    add_image_placeholder(
        doc,
        "商业化路径时间轴",
        "从“场地分成/核销”到“社区广告/赞助路线”再到“Echoland Night/城市周末/音乐节 IP”的三阶段路线图。",
    )

    add_heading(doc, "12. 运营与增长策略", 1)
    add_heading(doc, "12.1 首发城市与圈层选择", 2)
    add_para(
        doc,
        "建议以南京作为首发城市。南京有高校、青年文化空间、livehouse、独立书店、展览空间、夜游资源和城市漫游场景，适合用较轻的运营团队验证社区与线下组织力。早期不要覆盖全城所有活动，应先聚焦独立音乐、放映、市集、书店、摄影和夜间城市探索。",
    )
    add_heading(doc, "12.2 冷启动策略", 2)
    add_simple_table(
        doc,
        ["目标", "动作", "衡量指标"],
        [
            ["做出社区 vibe", "筛选 50-100 名种子用户，准备 30-50 条高质量种子帖。", "发帖率、评论率、关注率、次日留存。"],
            ["做出内容密度", "人工录入 50 条 PGC 活动和 80 条 UGC 回声点。", "收藏率、详情点击率、路线加入率。"],
            ["证明线下组织力", "举办 1 场 Echoland Night，目标 80-150 人。", "报名、到场、二次发帖、关注关系增长。"],
            ["证明场地价值", "访谈并合作 10-20 家小型场地。", "核销人数、老板复投意愿、分成收入。"],
        ],
        [1.1, 3.25, 1.95],
    )
    add_heading(doc, "12.3 双飞轮增长模型", 2)
    add_para(
        doc,
        "Echoland 应同时转动两个飞轮。第一个是社交 vibe 飞轮：用户因为社区气质进入、互动、发帖、关注同频的人，社区因此更鲜明。第二个是线下体验飞轮：用户因为无聊而发现活动或回声点，通过 AI 路线和邀约出门，线下体验回流成帖子、回声点和关系。",
    )
    add_image_placeholder(
        doc,
        "双飞轮增长模型",
        "两个相互连接的飞轮：社交 vibe 飞轮与线下体验飞轮。中间用“社区身份感”和“真实到场数据”连接。",
    )

    add_heading(doc, "13. 风险、商业漏洞与应对方案", 1)
    add_para(
        doc,
        "Echoland 的最大风险是同时想做太多：社交、活动、地图、AI、票务、线下活动和品牌广告。如果没有主轴，产品会被竞品拆分击穿。主轴必须保持清晰：有线下出口的氛围型城市社交社区。",
    )
    add_simple_table(
        doc,
        ["风险/漏洞", "严厉问题", "应对方案"],
        [
            ["定位过宽", "你到底是社交产品、活动平台、地图，还是 AI 工具？", "主定位锁定社交社区，活动/回声点/路线服务“打破无聊并真实出门”。"],
            ["冷启动困难", "没人发帖、没活动、没邀约，用户为什么留下？", "先做南京一个圈层，用编辑精选和线下联名场制造初始密度。"],
            ["社区无 vibe", "如何避免变成普通论坛或低配小红书？", "筛选首批用户，限制内容形态，建立发帖语气和视觉规范。"],
            ["UGC 安全风险", "用户上传废墟、天台、隐秘点位，出事怎么办？", "只允许公共合法低风险点位，做审核、举报、风险提示和失效反馈。"],
            ["场地不配合", "老板为什么还要多维护一个平台？", "不要求重写内容，支持推文/海报一键上传，并给曝光、收藏、核销数据。"],
            ["AI 变噱头", "用户需要 AI，还是需要推荐几个靠谱地方？", "AI 做决策压缩和路线执行，不做泛聊天。"],
            ["广告伤害社区", "社区商业化后会不会变味？", "广告必须原生、少量、符合城市生活方式调性。"],
            ["分成天花板低", "小场地单场收入有限，能养团队吗？", "早期验证到场，中期叠加赞助路线、品牌活动和自有 IP。"],
            ["城市扩张重运营", "南京跑通后，其他城市靠什么复制？", "沉淀城市编辑 SOP、场地合作模板、UGC 审核标准和活动数据结构。"],
        ],
        [1.15, 2.1, 3.05],
    )
    add_heading(doc, "13.1 安全与信任机制", 2)
    add_bullets(
        doc,
        [
            "邀约优先推荐公共场所集合，避免私密空间和高风险点位。",
            "邀约显示人数、状态、备注和风险提示，用户可举报、取消和屏蔽。",
            "UGC 回声点必须经过机器初筛与人工审核，危险点位不进入公开地图。",
            "用户可对点位标记“不安全”“已失效”“位置不准”。",
            "线下联名活动应有明确主办方、场地责任、活动规则和紧急联系信息。",
        ],
    )

    add_heading(doc, "14. 阶段里程碑与关键指标", 1)
    add_simple_table(
        doc,
        ["阶段", "时间", "关键任务", "关键指标"],
        [
            ["Demo 完成", "0-1 个月", "完成高保真前端 Demo：四 Tab、详情、AI 路线、社区、邀约、我的。", "演示路径顺畅，所有关键按钮有反馈。"],
            ["南京冷启动", "1-3 个月", "录入活动和回声点，招募种子用户，访谈场地。", "1000 名种子用户、100 条内容、10 家场地访谈。"],
            ["线下验证", "3-6 个月", "举办 Echoland Night 和小型联名活动。", "3 场活动、单场 80-150 人、到场率 60%+。"],
            ["商业化试跑", "6-9 个月", "核销分成、活动 Boost、赞助路线。", "月度场地收入、复投场地数、赞助合作数。"],
            ["城市复制", "9-18 个月", "扩展到杭州/上海/成都等城市，建立城市编辑 SOP。", "第二城市内容密度、城市留存、场地合作复制率。"],
            ["IP 化", "18 个月+", "Echoland Weekend、城市文化节和音乐节雏形。", "品牌赞助、票务收入、跨场地动员力。"],
        ],
        [0.95, 0.8, 3.1, 1.45],
    )
    add_heading(doc, "14.1 北极星指标", 2)
    add_para(
        doc,
        "建议北极星指标不要只看日活或发帖数，而要看“被 Echoland 促成的真实出门次数”。这个指标可以由加入路线、发起邀约、报名、核销和线下活动到场共同估算。它同时反映内容价值、社交价值和商业化潜力。",
    )
    add_simple_table(
        doc,
        ["指标类型", "指标"],
        [
            ["社区粘性", "次日留存、7 日留存、发帖率、评论率、关注关系数。"],
            ["内容价值", "活动详情点击率、收藏率、回声点上传数、编辑精选率。"],
            ["行动转化", "AI 路线生成次数、加入路线率、邀约发起率、报名/核销率。"],
            ["商业化", "场地合作数、单场到场人数、核销收入、广告/赞助收入。"],
            ["安全质量", "UGC 审核通过率、举报率、失效点位率、邀约投诉率。"],
        ],
        [1.25, 5.05],
    )

    doc.add_page_break()
    add_heading(doc, "15. 图片插入清单", 1)
    add_para(doc, "以下图片位置已在正文中用灰色占位框标出。建议后续设计或产品同学按说明补图。")
    add_simple_table(
        doc,
        ["序号", "位置", "建议插入图片"],
        [
            ["1", "封面", "Echoland App 界面叠加南京夜景的品牌主视觉。"],
            ["2", "执行摘要", "商业逻辑总览飞轮图。"],
            ["3", "用户旅程", "用户从无聊刷手机到线下出门的旅程图。"],
            ["4", "市场分析", "服务消费、演出市场和中小型现场音乐的漏斗图或柱状图。"],
            ["5", "产品体系", "四 Tab、活动详情、邀约和路线之间的信息架构图。"],
            ["6", "内容供给", "PGC 活动与 UGC 回声点的内容生产和审核流程图。"],
            ["7", "技术实现", "AI 抽取、打标、审核、路线生成和反馈回流架构图。"],
            ["8", "竞品分析", "竞品定位二维图。"],
            ["9", "商业模式", "商业化路径时间轴。"],
            ["10", "运营增长", "社交 vibe 飞轮与线下体验飞轮双飞轮图。"],
        ],
        [0.55, 1.25, 4.5],
    )

    add_heading(doc, "16. 资料来源", 1)
    add_source(doc, "人民日报：2025 年全国营业性演出票房收入超 616 亿元", "https://paper.people.com.cn/rmrb/pc/content/202601/14/content_30132174.html")
    add_source(doc, "新浪财经：中演协、灯塔联合发布《2025 现场音乐演出市场报告》", "https://finance.sina.com.cn/tech/roll/2026-04-23/doc-inhvnhra5045038.shtml")
    add_source(doc, "美团研究院：2025 年生活服务消费 9 大趋势洞察", "https://www.meituan.com/news/NN260121213001639")
    add_source(doc, "复旦发展研究院：2025 年轻人生活方式报告", "https://fddi.fudan.edu.cn/54/c7/c19047a742599/page.htm")
    add_source(doc, "秀动 ShowStart App Store 产品说明", "https://apps.apple.com/us/app/id923912459")
    add_source(doc, "豆瓣同城活动帮助中心", "https://help.douban.com/event?app=3")
    add_source(doc, "小红书从种草到生活兴趣社区研究报告", "https://pdf.dfcfw.com/pdf/H301_AP202510121760705963_1.pdf")
    add_source(doc, "大麦 App Store 产品说明", "https://apps.apple.com/cn/app/id513829338")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
